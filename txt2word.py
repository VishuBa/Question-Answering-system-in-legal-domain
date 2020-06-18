from docx import Document
import re
import os



path = r'C:\Users\prasad\Desktop\SKT_writ'
direct = os.listdir(path)
#onlyfiles = [f for f in listdir(mypath) if isfile(join(mypath, f))]
file = open(r'C:\Users\prasad\Desktop\SKT_writ\SKT_writ.txt')
def text2word(file):

	for i in direct:
	    document = Document()
	    #document.add_heading('SKT Towers Petetion', 0)
	    myfile = file.read()
	    myfile = re.sub(r'[^\x00-\x7F]+|\x0c',' ', myfile) # remove all non-XML-compatible characters
	    p = document.add_paragraph(myfile)
	    document.save()